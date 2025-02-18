import os
from flask import Flask, render_template, request, jsonify
from docx import Document
from datetime import datetime

app = Flask(__name__)
# Função para preencher a ATA com os dados fornecidos, incluindo tabelas
def preencher_ata(dados):
    # Obtém a data da reunião
    data_reuniao = dados.get("data", "")
    data_formatada = datetime.strptime(data_reuniao, "%Y-%m-%d")
    ano = data_formatada.strftime("%Y")
    mes = data_formatada.strftime("%m")

    # Cria diretório no formato "ata_reunioes/AAAA/MM"
    caminho_pasta = os.path.join("ata_reunioes", ano, mes)
    os.makedirs(caminho_pasta, exist_ok=True)

    # Nome do arquivo baseado na data
    nome_arquivo = f"ata_{data_formatada.strftime('%d-%m-%Y')}.docx"
    caminho_arquivo = os.path.join(caminho_pasta, nome_arquivo)

    # Preenche o template da ATA
    modelo_path = "template-ata.docx"
    doc = Document(modelo_path)

    # Substituir placeholders nos parágrafos
    for paragrafo in doc.paragraphs:
        for marcador, valor in dados.items():
            placeholder = f"{{{marcador}}}"  # Formato correto: {data}, {horaInicio}, etc.
            if placeholder in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(placeholder, valor)

    # Substituir placeholders nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for marcador, valor in dados.items():
                    placeholder = f"{{{marcador}}}"  # Formato correto: {data}, {horaInicio}, etc.
                    if placeholder in celula.text:
                        celula.text = celula.text.replace(placeholder, valor)

    # Salva a ata no diretório correto
    doc.save(caminho_arquivo)

    return caminho_arquivo  # Retorna o caminho onde foi salvo

# Rota para exibir a página HTML
@app.route('/')
def index():
    return render_template('index.html')

# Rota para processar os dados do formulário e gerar a ata
@app.route('/gerar_ata', methods=['POST'])
def gerar_ata():
    dados_reuniao = request.json  # Recebe os dados do formulário como JSON
    caminho_ata = preencher_ata(dados_reuniao)  # Cria a ata e obtém o caminho

    return jsonify({"mensagem": "ATA criada com sucesso!", "caminho": caminho_ata})

if __name__ == '__main__':
    app.run(debug=True)
